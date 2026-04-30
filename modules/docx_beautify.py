"""
docx_beautify.py — Reusable DOCX Beautification Module
========================================================
Drop-in module for creating professional Word documents with python-docx.
Provides presets, table formatting, Markdown conversion, and OXML helpers.

Dependencies: pip install python-docx

Usage:
    from docx_beautify import create_document, md_to_docx, add_professional_table

    # Quick preset document
    doc = create_document(preset="executive")

    # Markdown to DOCX
    md_to_docx("input.md", "output.docx", preset="technical")
"""

import re
import os
import json
from datetime import date
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ═══════════════════════════════════════════════════════════════════════════
# COLOR PALETTES
# ═══════════════════════════════════════════════════════════════════════════

PALETTES = {
    "fortive": {
        "primary":    "003366",  # Dark navy
        "secondary":  "005EB8",  # Corporate blue
        "accent":     "F3C13A",  # Gold
        "text":       "333333",  # Dark gray
        "light_text": "4A4A4A",  # Medium gray — darkened from 666666
        "bg_light":   "F7F7F7",  # Light gray background
        "bg_alt":     "F2F6FA",  # Alternating row — lighter than EDF2F7
        "header_bg":  "1A4A6E",  # Softer header — less heavy than 003366
        "header_fg":  "FFFFFF",  # Table header text
        "border":     "CCCCCC",  # Border gray
        "success":    "28A745",  # Green
        "warning":    "F3C13A",  # Gold/yellow
        "danger":     "DC3545",  # Red
        "info":       "17A2B8",  # Teal
        "code_bg":    "F5F5F5",  # Code block background
    },
    "executive": {
        "primary":    "1B3A5C",  # Deep navy (heading 1, bold keywords primary)
        "secondary":  "3D556F",  # Softer blue-gray (heading 2, bold keywords secondary)
        "accent":     "C5A55A",  # Muted gold
        "text":       "2D2D2D",  # Body text (softer than pure black)
        "light_text": "555555",  # Footers, muted labels — darkened from 717171
        "bg_light":   "FAFAFA",
        "bg_alt":     "F5F8FB",  # Lighter alternating row — less heavy than F0F4F8
        "header_bg":  "2C4F6E",  # Softer header — less heavy than pure 1B3A5C
        "header_fg":  "FFFFFF",
        "border":     "D0D0D0",
        "success":    "2E7D32",
        "warning":    "ED6C02",
        "danger":     "C62828",
        "info":       "0277BD",
        "code_bg":    "F8F8F8",
    },
    "modern": {
        "primary":    "4A90D9",  # Bright blue
        "secondary":  "6C5CE7",  # Purple
        "accent":     "00B894",  # Emerald
        "text":       "2C3E50",
        "light_text": "7F8C8D",
        "bg_light":   "F8F9FA",
        "bg_alt":     "E8F4FD",
        "header_bg":  "4A90D9",
        "header_fg":  "FFFFFF",
        "border":     "DEE2E6",
        "success":    "00B894",
        "warning":    "FDCB6E",
        "danger":     "E17055",
        "info":       "74B9FF",
        "code_bg":    "F1F3F5",
    },
    "minimal": {
        "primary":    "333333",
        "secondary":  "555555",
        "accent":     "0066CC",
        "text":       "333333",
        "light_text": "888888",
        "bg_light":   "FAFAFA",
        "bg_alt":     "F5F5F5",
        "header_bg":  "333333",
        "header_fg":  "FFFFFF",
        "border":     "E0E0E0",
        "success":    "4CAF50",
        "warning":    "FF9800",
        "danger":     "F44336",
        "info":       "2196F3",
        "code_bg":    "F5F5F5",
    },
}


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT PRESETS
# ═══════════════════════════════════════════════════════════════════════════

PRESETS = {
    "executive": {
        "palette": "executive",
        "font_body": "Calibri",
        "font_heading": "Calibri",
        "font_code": "Consolas",
        "size_body": 10.5,
        "size_h1": 13,
        "size_h2": 12,
        "size_h3": 11,
        "size_h4": 10.5,
        "size_code": 8.5,
        "size_table": 9,
        "size_footer": 8,
        "margin_top": 2.54,     # cm (~1 inch)
        "margin_bottom": 2.54,
        "margin_left": 2.54,
        "margin_right": 2.54,
        "line_spacing": 1.15,
        "space_after_para": 6,  # pt — tight like MAQ SOW
        "heading_color": False,  # Black headings — professional restraint
        "heading_h1_color": True,  # Allow subtle color on H1 only
        "table_style": "professional",
    },
    "technical": {
        "palette": "fortive",
        "font_body": "Calibri",
        "font_heading": "Calibri",
        "font_code": "Consolas",
        "size_body": 10,
        "size_h1": 13,
        "size_h2": 12,
        "size_h3": 11,
        "size_h4": 10,
        "size_code": 8.5,
        "size_table": 9,
        "size_footer": 8,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_left": 2.54,
        "margin_right": 2.54,
        "line_spacing": 1.15,
        "space_after_para": 6,
        "heading_color": False,
        "heading_h1_color": True,
        "table_style": "professional",
    },
    "memo": {
        "palette": "minimal",
        "font_body": "Calibri",
        "font_heading": "Calibri",
        "font_code": "Consolas",
        "size_body": 10.5,
        "size_h1": 12,
        "size_h2": 11,
        "size_h3": 10.5,
        "size_h4": 10.5,
        "size_code": 9,
        "size_table": 9,
        "size_footer": 8,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_left": 2.54,
        "margin_right": 2.54,
        "line_spacing": 1.15,
        "space_after_para": 6,
        "heading_color": False,
        "heading_h1_color": False,
        "table_style": "simple",
    },
    "report": {
        "palette": "fortive",
        "font_body": "Calibri",
        "font_heading": "Calibri",
        "font_code": "Consolas",
        "size_body": 10.5,
        "size_h1": 14,
        "size_h2": 12,
        "size_h3": 11,
        "size_h4": 10.5,
        "size_code": 9,
        "size_table": 9,
        "size_footer": 8,
        "margin_top": 2.54,
        "margin_bottom": 2.54,
        "margin_left": 2.54,
        "margin_right": 2.54,
        "line_spacing": 1.15,
        "space_after_para": 6,
        "heading_color": False,
        "heading_h1_color": True,
        "table_style": "professional",
    },
}


# ═══════════════════════════════════════════════════════════════════════════
# OXML HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color):
    """Set table cell background color via OXML."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = tc_pr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    # Remove existing shading
    for existing in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(existing)
    tc_pr.append(shading)


def set_paragraph_shading(paragraph, color):
    """Set paragraph background color via OXML (for callout boxes, code blocks)."""
    pPr = paragraph._element.get_or_add_pPr()
    shd = pPr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    pPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders. Each border is a dict: {color, size, style}."""
    tc_pr = cell._element.get_or_add_tcPr()
    borders = tc_pr.makeelement(qn("w:tcBorders"), {})
    for existing in tc_pr.findall(qn("w:tcBorders")):
        tc_pr.remove(existing)
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            border = borders.makeelement(qn(f"w:{side}"), {
                qn("w:val"): val.get("style", "single"),
                qn("w:sz"): str(val.get("size", 4)),
                qn("w:color"): val.get("color", "000000"),
                qn("w:space"): "0",
            })
            borders.append(border)
    tc_pr.append(borders)


def set_cell_margins(cell, top=0, bottom=0, left=80, right=80):
    """Set internal cell margins (padding) in twips."""
    tc_pr = cell._element.get_or_add_tcPr()
    margins = tc_pr.makeelement(qn("w:tcMar"), {})
    for existing in tc_pr.findall(qn("w:tcMar")):
        tc_pr.remove(existing)
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        m = margins.makeelement(qn(f"w:{side}"), {
            qn("w:w"): str(val),
            qn("w:type"): "dxa",
        })
        margins.append(m)
    tc_pr.append(margins)


def set_cell_vertical_alignment(cell, align="center"):
    """Set vertical alignment: top, center, bottom."""
    tc_pr = cell._element.get_or_add_tcPr()
    val_el = tc_pr.makeelement(qn("w:vAlign"), {qn("w:val"): align})
    for existing in tc_pr.findall(qn("w:vAlign")):
        tc_pr.remove(existing)
    tc_pr.append(val_el)


def add_page_number(paragraph):
    """Add a PAGE field code to a paragraph (for headers/footers)."""
    run = paragraph.add_run()
    fld_char_begin = run._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = run2._element.makeelement(qn("w:instrText"), {qn("xml:space"): "preserve"})
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = run3._element.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
    run3._element.append(fld_char_end)


def add_horizontal_rule(doc, color="CCCCCC", thickness=1):
    """Add a thin horizontal rule as a paragraph bottom border."""
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    borders = pPr.makeelement(qn("w:pBdr"), {})
    bottom = borders.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): str(thickness * 4),
        qn("w:space"): "1",
        qn("w:color"): color,
    })
    borders.append(bottom)
    pPr.append(borders)
    return p


def set_no_spacing(paragraph):
    """Remove before/after spacing from a paragraph."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ═══════════════════════════════════════════════════════════════════════════

def create_document(preset="executive", palette_override=None):
    """Create a new Document with preset styles applied."""
    cfg = PRESETS.get(preset, PRESETS["executive"]).copy()
    pal = PALETTES.get(palette_override or cfg["palette"], PALETTES["fortive"]).copy()

    doc = Document()
    _apply_preset(doc, cfg, pal)
    return doc


def _apply_preset(doc, cfg, pal):
    """Apply a preset configuration to a document."""
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(cfg["margin_top"])
        section.bottom_margin = Cm(cfg["margin_bottom"])
        section.left_margin = Cm(cfg["margin_left"])
        section.right_margin = Cm(cfg["margin_right"])

    # Normal style
    style = doc.styles["Normal"]
    style.font.name = cfg["font_body"]
    style.font.size = Pt(cfg["size_body"])
    style.font.color.rgb = RGBColor.from_string(pal["text"])
    pf = style.paragraph_format
    pf.space_after = Pt(cfg["space_after_para"])
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = cfg["line_spacing"]

    # Heading styles — MAQ-inspired: tight spacing, minimal size jumps, restrained color
    use_all_heading_color = cfg.get("heading_color", False)
    use_h1_color_only = cfg.get("heading_h1_color", False)
    heading_text_color = RGBColor.from_string(pal["text"])  # default: near-black
    h1_color = RGBColor.from_string(pal["primary"]) if (use_all_heading_color or use_h1_color_only) else heading_text_color

    for level, size_key in [(1, "size_h1"), (2, "size_h2"), (3, "size_h3"), (4, "size_h4")]:
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h = doc.styles[style_name]
            h.font.name = cfg["font_heading"]
            h.font.size = Pt(cfg[size_key])
            h.font.bold = True
            # Color: only H1 gets subtle color (if h1_color_only), rest are dark text
            if use_all_heading_color:
                h.font.color.rgb = RGBColor.from_string(pal["primary"])
            elif level == 1 and use_h1_color_only:
                h.font.color.rgb = h1_color
            else:
                h.font.color.rgb = heading_text_color
            hpf = h.paragraph_format
            # Tight spacing: 12pt before H1, 10pt before H2, 8pt before H3/H4
            hpf.space_before = Pt(12 if level == 1 else 10 if level == 2 else 8)
            hpf.space_after = Pt(4 if level <= 2 else 3)
            hpf.keep_with_next = True

    # List styles
    for list_style in ["List Bullet", "List Number"]:
        if list_style in doc.styles:
            ls = doc.styles[list_style]
            ls.font.name = cfg["font_body"]
            ls.font.size = Pt(cfg["size_body"])


def apply_preset_to_existing(doc, preset="executive", palette_override=None):
    """Apply preset styles to an existing document (for post-processing)."""
    cfg = PRESETS.get(preset, PRESETS["executive"]).copy()
    pal = PALETTES.get(palette_override or cfg["palette"], PALETTES["fortive"]).copy()
    _apply_preset(doc, cfg, pal)
    return doc


# ═══════════════════════════════════════════════════════════════════════════
# TABLE FORMATTING
# ═══════════════════════════════════════════════════════════════════════════

def add_professional_table(doc, headers, rows, palette="fortive", auto_size=True,
                           alt_rows=True, compact=False):
    """
    Add a professionally formatted table.

    Args:
        doc: Document object
        headers: list of header strings
        rows: list of lists (each inner list = one row of cell values)
        palette: color palette name
        auto_size: auto-calculate column widths from content
        alt_rows: alternating row shading
        compact: reduce cell padding for dense tables
    """
    pal = PALETTES.get(palette, PALETTES["fortive"])
    n_cols = len(headers)
    n_rows = len(rows)

    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    pad = (40, 40, 60, 60) if compact else (60, 60, 100, 100)

    # Header row
    for ci, header in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(header))
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(pal["header_fg"])
        run.font.name = "Calibri"
        set_cell_shading(cell, pal["header_bg"])
        set_cell_margins(cell, *pad)
        set_cell_vertical_alignment(cell, "center")

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, value in enumerate(row_data):
            if ci >= n_cols:
                break
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor.from_string(pal["text"])
            set_cell_margins(cell, *pad)
            set_cell_vertical_alignment(cell, "center")

            if alt_rows and ri % 2 == 1:
                set_cell_shading(cell, pal["bg_alt"])

    # Auto-size columns based on content
    if auto_size:
        _auto_size_columns(table, headers, rows)

    return table


def add_key_value_table(doc, pairs, palette="fortive", key_width_pct=30):
    """
    Add a two-column key-value table (for metadata, parameters, etc).

    Args:
        pairs: list of (key, value) tuples
        key_width_pct: percentage of table width for the key column
    """
    pal = PALETTES.get(palette, PALETTES["fortive"])
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, (key, value) in enumerate(pairs):
        # Key cell
        k_cell = table.rows[ri].cells[0]
        k_cell.text = ""
        kp = k_cell.paragraphs[0]
        kr = kp.add_run(str(key))
        kr.font.bold = True
        kr.font.size = Pt(9)
        kr.font.name = "Calibri"
        kr.font.color.rgb = RGBColor.from_string(pal["text"])
        set_cell_shading(k_cell, pal["bg_light"])
        set_cell_margins(k_cell, 60, 60, 100, 100)

        # Value cell
        v_cell = table.rows[ri].cells[1]
        v_cell.text = ""
        vp = v_cell.paragraphs[0]
        vr = vp.add_run(str(value))
        vr.font.size = Pt(9)
        vr.font.name = "Calibri"
        vr.font.color.rgb = RGBColor.from_string(pal["text"])
        set_cell_margins(v_cell, 60, 60, 100, 100)

        if ri % 2 == 1:
            set_cell_shading(v_cell, pal["bg_alt"])

    return table


def add_status_table(doc, items, palette="fortive"):
    """
    Add a table with status badges (colored dots + text).

    Args:
        items: list of dicts with keys: name, status, detail (optional)
               status values: "pass"/"success"/"complete", "fail"/"error", "warn"/"warning", "info"/"pending"
    """
    pal = PALETTES.get(palette, PALETTES["fortive"])
    status_colors = {
        "pass": pal["success"], "success": pal["success"], "complete": pal["success"],
        "fail": pal["danger"], "error": pal["danger"],
        "warn": pal["warning"], "warning": pal["warning"],
        "info": pal["info"], "pending": pal["info"],
    }
    status_symbols = {
        "pass": "\u2713", "success": "\u2713", "complete": "\u2713",  # checkmark
        "fail": "\u2717", "error": "\u2717",  # X
        "warn": "\u26A0", "warning": "\u26A0",  # warning triangle
        "info": "\u2139", "pending": "\u2139",  # info
    }

    has_detail = any(item.get("detail") for item in items)
    headers = ["Item", "Status"] + (["Detail"] if has_detail else [])
    rows = []
    for item in items:
        status_key = item.get("status", "info").lower()
        row = [item["name"], status_key.upper()]
        if has_detail:
            row.append(item.get("detail", ""))
        rows.append(row)

    table = add_professional_table(doc, headers, rows, palette=palette)

    # Apply status coloring to the status column
    for ri, item in enumerate(items):
        status_key = item.get("status", "info").lower()
        color = status_colors.get(status_key, pal["info"])
        symbol = status_symbols.get(status_key, "\u2022")

        cell = table.rows[ri + 1].cells[1]
        cell.text = ""
        p = cell.paragraphs[0]

        # Symbol run
        sym_run = p.add_run(symbol + " ")
        sym_run.font.color.rgb = RGBColor.from_string(color)
        sym_run.font.size = Pt(10)
        sym_run.font.bold = True

        # Status text
        txt_run = p.add_run(status_key.upper())
        txt_run.font.size = Pt(9)
        txt_run.font.bold = True
        txt_run.font.color.rgb = RGBColor.from_string(color)

    return table


def _auto_size_columns(table, headers, rows):
    """Auto-size table columns based on content length (heuristic)."""
    n_cols = len(headers)
    max_lens = [len(str(h)) for h in headers]
    for row in rows:
        for ci in range(min(len(row), n_cols)):
            max_lens[ci] = max(max_lens[ci], len(str(row[ci])))

    # Clamp: min 8 chars, max 60 chars
    max_lens = [max(8, min(60, ml)) for ml in max_lens]
    total = sum(max_lens)

    # Available width: US Letter with 2.5cm margins ~ 16cm ~ 9070 twips
    avail = 9070
    for ci in range(n_cols):
        w = int(avail * max_lens[ci] / total)
        for row in table.rows:
            row.cells[ci].width = Emu(w * 635)  # twips to EMU


# ═══════════════════════════════════════════════════════════════════════════
# CONTENT HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def add_code_block(doc, text, palette="fortive", language=None):
    """Add a monospace code block with background shading."""
    pal = PALETTES.get(palette, PALETTES["fortive"])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)

    set_paragraph_shading(p, pal["code_bg"])

    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(pal["text"])
    return p


def add_callout_box(doc, text, style="info", palette="fortive"):
    """
    Add a colored callout box (like a GitHub alert).

    Args:
        style: "info", "warning", "success", "danger"
    """
    pal = PALETTES.get(palette, PALETTES["fortive"])
    color_map = {
        "info": pal["info"],
        "warning": pal["warning"],
        "success": pal["success"],
        "danger": pal["danger"],
    }
    bg_map = {
        "info": "E8F4FD",
        "warning": "FFF8E1",
        "success": "E8F5E9",
        "danger": "FFEBEE",
    }
    icon_map = {
        "info": "\u2139\uFE0F",
        "warning": "\u26A0\uFE0F",
        "success": "\u2705",
        "danger": "\u274C",
    }

    color = color_map.get(style, pal["info"])
    bg = bg_map.get(style, "E8F4FD")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)

    set_paragraph_shading(p, bg)

    # Left border via OXML
    pPr = p._element.get_or_add_pPr()
    borders = pPr.makeelement(qn("w:pBdr"), {})
    left_border = borders.makeelement(qn("w:left"), {
        qn("w:val"): "single",
        qn("w:sz"): "24",
        qn("w:space"): "8",
        qn("w:color"): color,
    })
    borders.append(left_border)
    pPr.append(borders)

    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = "Calibri"
    return p


def add_cover_page(doc, title, subtitle=None, author=None, date_str=None,
                   palette="fortive"):
    """Add a simple cover page with title, subtitle, author, and date."""
    pal = PALETTES.get(palette, PALETTES["fortive"])
    date_str = date_str or date.today().strftime("%B %d, %Y")

    # Spacer — moderate, not excessive
    for _ in range(5):
        p = doc.add_paragraph()
        set_no_spacing(p)

    # Title — restrained: 18pt bold, not giant 36pt
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(pal["primary"])
    run.font.name = "Calibri"

    # Thin accent line
    add_horizontal_rule(doc, color=pal["border"], thickness=1)

    # Subtitle
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string(pal["light_text"])
        run.font.name = "Calibri"

    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        set_no_spacing(p)

    # Author
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(author)
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string(pal["text"])

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(date_str)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(pal["light_text"])

    # Page break
    doc.add_page_break()


def add_header_footer(doc, header_text=None, footer_text=None, page_numbers=True,
                      palette="fortive"):
    """Add header and/or footer to all sections."""
    pal = PALETTES.get(palette, PALETTES["fortive"])

    for section in doc.sections:
        if header_text:
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            p.text = ""
            run = p.add_run(header_text)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(pal["light_text"])
            run.font.name = "Calibri"
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Bottom border on header
            pPr = p._element.get_or_add_pPr()
            borders = pPr.makeelement(qn("w:pBdr"), {})
            bottom = borders.makeelement(qn("w:bottom"), {
                qn("w:val"): "single",
                qn("w:sz"): "4",
                qn("w:space"): "4",
                qn("w:color"): pal["border"],
            })
            borders.append(bottom)
            pPr.append(borders)

        if footer_text or page_numbers:
            footer = section.footer
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if footer_text:
                run = p.add_run(footer_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor.from_string(pal["light_text"])
                run.font.name = "Calibri"

            if page_numbers:
                if footer_text:
                    run = p.add_run("  |  Page ")
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor.from_string(pal["light_text"])
                else:
                    run = p.add_run("Page ")
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor.from_string(pal["light_text"])
                add_page_number(p)


# ═══════════════════════════════════════════════════════════════════════════
# INLINE TEXT FORMATTING
# ═══════════════════════════════════════════════════════════════════════════

def add_formatted_text(paragraph, text, palette="fortive"):
    """Parse inline Markdown formatting: **bold**, `code`, *italic*, [links]."""
    pal = PALETTES.get(palette, PALETTES["fortive"])
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(pal["light_text"])
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ═══════════════════════════════════════════════════════════════════════════
# MARKDOWN → DOCX ENGINE
# ═══════════════════════════════════════════════════════════════════════════

def md_to_docx(md_path, docx_path, preset="executive", title=None, author=None,
               cover_page=False, header_footer=True):
    """
    Convert a Markdown file to a professionally formatted DOCX.

    Args:
        md_path: path to input .md file
        docx_path: path to output .docx file
        preset: document preset name
        title: document title (for cover page and header)
        author: author name (for cover page)
        cover_page: whether to add a cover page
        header_footer: whether to add header/footer with page numbers
    """
    cfg = PRESETS.get(preset, PRESETS["executive"])
    pal_name = cfg["palette"]

    doc = create_document(preset=preset)

    # Extract title from first H1 if not provided
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    if not title:
        m = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        if m:
            title = m.group(1).strip()

    # Cover page
    if cover_page and title:
        add_cover_page(doc, title, author=author, palette=pal_name)

    # Parse and add content
    _parse_markdown(doc, content, cfg, pal_name)

    # Header/footer
    if header_footer:
        add_header_footer(
            doc,
            header_text=title or os.path.basename(md_path),
            page_numbers=True,
            palette=pal_name,
        )

    doc.save(docx_path)
    size_kb = os.path.getsize(docx_path) / 1024
    print(f"Saved: {docx_path} ({size_kb:.0f} KB)")
    return docx_path


def _parse_markdown(doc, content, cfg, palette):
    """Parse Markdown content and add to document."""
    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # --- Code blocks ---
        if stripped.lstrip().startswith("```"):
            if in_code_block:
                add_code_block(doc, "\n".join(code_lines), palette=palette)
                code_lines = []
                in_code_block = False
            else:
                _flush_table(doc, table_lines, in_table, palette)
                table_lines = []
                in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # --- Tables ---
        if stripped.startswith("|"):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue
        elif in_table:
            _flush_table(doc, table_lines, in_table, palette)
            table_lines = []
            in_table = False

        # --- HTML tags (skip) ---
        if stripped.startswith("<details") or stripped.startswith("</details"):
            i += 1
            continue
        if stripped.startswith("<summary"):
            m = re.search(r"<summary>(.*?)</summary>", stripped)
            if m:
                p = doc.add_paragraph()
                run = p.add_run(m.group(1))
                run.italic = True
            i += 1
            continue

        # --- Headings ---
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Strip inline formatting from heading text
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # --- Horizontal rule ---
        if stripped in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # --- Empty line ---
        if not stripped.strip():
            i += 1
            continue

        # --- Bullet points ---
        bullet_match = re.match(r"^(\s*)[-*]\s+(.+)", stripped)
        if bullet_match:
            indent = len(bullet_match.group(1))
            text = bullet_match.group(2)
            # Handle checkboxes
            text = text.replace("[ ]", "\u2610").replace("[x]", "\u2611").replace("[X]", "\u2611")
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 4:
                p.paragraph_format.left_indent = Cm(1.5)
            add_formatted_text(p, text, palette=palette)
            i += 1
            continue

        # --- Numbered lists ---
        num_match = re.match(r"^(\d+)\.\s+(.+)", stripped.strip())
        if num_match:
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, num_match.group(2), palette=palette)
            i += 1
            continue

        # --- Block quotes ---
        if stripped.startswith("> "):
            text = stripped[2:].strip()
            add_callout_box(doc, text, style="info", palette=palette)
            i += 1
            continue

        # --- Regular paragraph ---
        p = doc.add_paragraph()
        add_formatted_text(p, stripped.strip(), palette=palette)
        i += 1

    # Flush remaining
    if in_code_block and code_lines:
        add_code_block(doc, "\n".join(code_lines), palette=palette)
    _flush_table(doc, table_lines, in_table, palette)


def _flush_table(doc, table_lines, in_table, palette):
    """Flush accumulated table lines into a formatted table."""
    if not in_table or not table_lines:
        return

    # Parse markdown table
    header_line = table_lines[0]
    data_lines = []
    for tl in table_lines[1:]:
        # Skip separator rows
        cells = [c.strip() for c in tl.split("|")[1:-1]]
        if all(re.match(r"^[-:]+$", c) for c in cells if c.strip()):
            continue
        data_lines.append(tl)

    headers = [c.strip().strip("*") for c in header_line.split("|")[1:-1]]
    headers = [h for h in headers if h]
    if not headers:
        return

    rows = []
    for dl in data_lines:
        cells = [c.strip() for c in dl.split("|")[1:-1]]
        rows.append(cells)

    add_professional_table(doc, headers, rows, palette=palette)
    doc.add_paragraph()  # Spacing after table


# ═══════════════════════════════════════════════════════════════════════════
# CONVENIENCE: Beautify an existing DOCX
# ═══════════════════════════════════════════════════════════════════════════

def beautify_tables(doc, palette="fortive", alt_rows=True):
    """Re-format all tables in an existing document with professional styling."""
    pal = PALETTES.get(palette, PALETTES["fortive"])

    for table in doc.tables:
        if not table.rows:
            continue

        # Style header row
        for cell in table.rows[0].cells:
            set_cell_shading(cell, pal["header_bg"])
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor.from_string(pal["header_fg"])
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
            set_cell_margins(cell, 60, 60, 100, 100)

        # Style data rows
        for ri in range(1, len(table.rows)):
            for cell in table.rows[ri].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
                set_cell_margins(cell, 60, 60, 100, 100)
                if alt_rows and ri % 2 == 0:
                    set_cell_shading(cell, pal["bg_alt"])


# ═══════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (via OXML field codes — borrowed from docXMLater pattern)
# ═══════════════════════════════════════════════════════════════════════════

def add_toc(doc, title="Table of Contents", levels=3, tab_leader="dot",
            after_paragraph=None):
    """Add a Table of Contents that Word populates on open/update.

    Uses SDT (Structured Document Tag) wrapping ECMA-376 complex field codes.
    The TOC appears as a placeholder until the user right-clicks → Update Field
    or presses Ctrl+A → F9 in Word.

    Args:
        doc: Document object
        title: TOC heading text (None to omit heading)
        levels: Heading levels to include (1-9, default 3)
        tab_leader: Leader between entry and page number
                    ("dot", "hyphen", "underscore", "none")
        after_paragraph: Insert TOC after this paragraph element.
                         If None, appends to end of document body.

    Returns:
        The SDT OxmlElement (for positioning control).

    Per ECMA-376 §17.16.5: All complex fields MUST have
    begin → instrText → separate → content → end structure.
    Missing any element causes Word to reject the document as corrupted.
    """
    from docx.oxml import OxmlElement

    # Build field instruction
    instruction = f'TOC \\o "1-{levels}" \\h \\z'
    leader_map = {"hyphen": "h", "underscore": "u", "none": "n"}
    if tab_leader in leader_map:
        instruction += f' \\p "{leader_map[tab_leader]}"'
    instruction += " \\* MERGEFORMAT"

    # SDT wrapper
    sdt = OxmlElement("w:sdt")
    sdtPr = OxmlElement("w:sdtPr")
    docPartObj = OxmlElement("w:docPartObj")
    docPartGallery = OxmlElement("w:docPartGallery")
    docPartGallery.set(qn("w:val"), "Table of Contents")
    docPartUnique = OxmlElement("w:docPartUnique")
    docPartObj.append(docPartGallery)
    docPartObj.append(docPartUnique)
    sdtPr.append(docPartObj)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")

    # Title paragraph (styled as TOCHeading)
    if title:
        p_title = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), "TOCHeading")
        pPr.append(pStyle)
        p_title.append(pPr)
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = title
        r.append(t)
        p_title.append(r)
        sdtContent.append(p_title)

    # Field paragraph: begin → instrText → separate → placeholder → end
    p = OxmlElement("w:p")

    # 1. begin
    r1 = OxmlElement("w:r")
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    r1.append(fldBegin)
    p.append(r1)

    # 2. instrText
    r2 = OxmlElement("w:r")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instruction
    r2.append(instrText)
    p.append(r2)

    # 3. separate
    r3 = OxmlElement("w:r")
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    r3.append(fldSep)
    p.append(r3)

    # 4. placeholder content
    r4 = OxmlElement("w:r")
    rPr4 = OxmlElement("w:rPr")
    noProof = OxmlElement("w:noProof")
    rPr4.append(noProof)
    r4.append(rPr4)
    t4 = OxmlElement("w:t")
    t4.text = "Right-click to update field."
    r4.append(t4)
    p.append(r4)

    # 5. end
    r5 = OxmlElement("w:r")
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    r5.append(fldEnd)
    p.append(r5)

    sdtContent.append(p)
    sdt.append(sdtContent)

    # Insert position
    if after_paragraph is not None:
        after_paragraph.addnext(sdt)
    else:
        doc.element.body.append(sdt)

    return sdt


# ═══════════════════════════════════════════════════════════════════════════
# MERMAID DIAGRAM SUPPORT
# Primary: beautiful-mermaid (npm) — zero-DOM SVG rendering, 15 themes
# Fallback 1: mmdc (mermaid-cli) — headless Puppeteer rendering
# Fallback 2: ASCII art via beautiful-mermaid
# Fallback 3: Styled code block
# ═══════════════════════════════════════════════════════════════════════════

# Path to built beautiful-mermaid dist (set after npm install)
_BEAUTIFUL_MERMAID_DIST = None

# beautiful-mermaid theme presets (bg/fg/accent/line/muted)
MERMAID_THEMES = {
    "corporate":  {"bg": "#FFFFFF", "fg": "#333333", "accent": "#005EB8",
                   "line": "#666666", "muted": "#999999",
                   "surface": "#F5F5F5", "border": "#CCCCCC"},
    "github-light": {"bg": "#ffffff", "fg": "#1f2328", "accent": "#0969da",
                     "line": "#d1d9e0", "muted": "#59636e"},
    "github-dark":  {"bg": "#0d1117", "fg": "#e6edf3", "accent": "#4493f8",
                     "line": "#3d444d", "muted": "#9198a1"},
    "tokyo-night":  {"bg": "#1a1b26", "fg": "#a9b1d6", "accent": "#7aa2f7",
                     "line": "#3d59a1", "muted": "#565f89"},
    "nord":         {"bg": "#2e3440", "fg": "#d8dee9", "accent": "#88c0d0",
                     "line": "#4c566a", "muted": "#616e88"},
    "clean":        {"bg": "#FFFFFF", "fg": "#27272A"},  # minimal mono
}


def _find_beautiful_mermaid():
    """Locate built beautiful-mermaid dist/index.js. Returns path or None."""
    global _BEAUTIFUL_MERMAID_DIST
    if _BEAUTIFUL_MERMAID_DIST is not None:
        return _BEAUTIFUL_MERMAID_DIST

    import shutil
    candidates = [
        os.path.join(os.path.dirname(__file__), "node_modules",
                     "beautiful-mermaid", "dist", "index.js"),
        "/tmp/beautiful-mermaid/dist/index.js",
        os.path.expanduser("~/AppData/Local/Temp/beautiful-mermaid/dist/index.js"),
        # Global npm install locations (admin and user accounts)
        os.path.expanduser("~/AppData/Roaming/npm/node_modules/beautiful-mermaid/dist/index.js"),
        r"<ADMIN_HOME>/AppData\Roaming\npm\node_modules\beautiful-mermaid\dist\index.js",
        r"<USER_HOME>/AppData\Roaming\npm\node_modules\beautiful-mermaid\dist\index.js",
    ]
    for p in candidates:
        if os.path.isfile(p):
            _BEAUTIFUL_MERMAID_DIST = p
            return p

    # Check if node is available for fallback
    if shutil.which("node") is None:
        return None
    return None


def _check_mmdc():
    """Check if mermaid-cli is available (global mmdc or via npx)."""
    import shutil
    if shutil.which("mmdc") is not None:
        return "mmdc"
    if shutil.which("npx") is not None:
        return "npx"
    return None


def render_mermaid_svg(mermaid_code, theme="corporate"):
    """Render mermaid code to SVG string using beautiful-mermaid (Node.js).

    This is the preferred rendering method — zero DOM, synchronous, themed.
    Supports: flowchart, sequence, class, ER, state, and XY chart diagrams.

    Args:
        mermaid_code: Mermaid diagram source code
        theme: Theme name from MERMAID_THEMES or a dict with bg/fg/accent keys

    Returns:
        SVG string, or None if beautiful-mermaid is not available.
    """
    import subprocess
    import tempfile

    dist_path = _find_beautiful_mermaid()
    if dist_path is None:
        return None

    colors = theme if isinstance(theme, dict) else MERMAID_THEMES.get(theme, MERMAID_THEMES["corporate"])

    # Build a small Node.js script that renders and prints SVG
    opts_json = json.dumps(colors)
    escaped_code = json.dumps(mermaid_code)

    # Use file:// URL for Windows compatibility with ESM imports
    file_url = "file:///" + dist_path.replace("\\", "/")
    script = (
        f'import {{ renderMermaidSVG }} from "{file_url}";\n'
        f'const svg = renderMermaidSVG({escaped_code}, {opts_json});\n'
        f'process.stdout.write(svg);\n'
    )

    with tempfile.NamedTemporaryFile(mode='w', suffix='.mjs', delete=False,
                                      encoding='utf-8') as f:
        f.write(script)
        script_path = f.name

    try:
        result = subprocess.run(
            ['node', script_path],
            capture_output=True, text=True, timeout=15
        )
        if result.returncode == 0 and result.stdout.startswith("<svg"):
            return result.stdout
        return None
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return None
    finally:
        try:
            os.unlink(script_path)
        except OSError:
            pass


def render_mermaid_ascii(mermaid_code):
    """Render mermaid code to Unicode box-drawing ASCII art.

    Uses beautiful-mermaid's ASCII renderer. Great for terminal output
    or embedding as monospace text in documents.

    Args:
        mermaid_code: Mermaid diagram source code

    Returns:
        ASCII string, or None if beautiful-mermaid is not available.
    """
    import subprocess
    import tempfile

    dist_path = _find_beautiful_mermaid()
    if dist_path is None:
        return None

    escaped_code = json.dumps(mermaid_code)
    file_url = "file:///" + dist_path.replace("\\", "/")
    script = (
        f'import {{ renderMermaidASCII }} from "{file_url}";\n'
        f'const ascii = renderMermaidASCII({escaped_code});\n'
        f'process.stdout.write(ascii);\n'
    )

    with tempfile.NamedTemporaryFile(mode='w', suffix='.mjs', delete=False,
                                      encoding='utf-8') as f:
        f.write(script)
        script_path = f.name

    try:
        result = subprocess.run(
            ['node', script_path],
            capture_output=True, text=True, timeout=15,
            encoding='utf-8', errors='replace'
        )
        if result.returncode == 0 and len(result.stdout) > 0:
            return result.stdout
        return None
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return None
    finally:
        try:
            os.unlink(script_path)
        except OSError:
            pass


def render_mermaid(mermaid_code, output_format="png", width=800, theme="default"):
    """Render mermaid code to an image file using mmdc CLI or npx.

    Tries global mmdc first, then npx @mermaid-js/mermaid-cli.

    Args:
        mermaid_code: Mermaid diagram source code
        output_format: "png" or "svg"
        width: Image width in pixels
        theme: Mermaid theme ("default", "dark", "forest", "neutral")

    Returns:
        Path to the rendered image file (caller must delete after use).

    Raises:
        RuntimeError: If neither mmdc nor npx is available, or rendering fails.
    """
    import subprocess
    import tempfile

    tool = _check_mmdc()
    if not tool:
        raise RuntimeError(
            "mermaid-cli not found. Install with: npm install -g @mermaid-js/mermaid-cli"
        )

    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False,
                                      encoding='utf-8') as f:
        f.write(mermaid_code)
        input_path = f.name

    output_path = input_path.replace('.mmd', f'.{output_format}')

    if tool == "mmdc":
        cmd = ['mmdc', '-i', input_path, '-o', output_path,
               '-b', 'transparent', '-w', str(width), '-t', theme]
        use_shell = False
    else:
        # npx — on Windows, npx is a .cmd file requiring shell=True
        cmd = f'npx -y @mermaid-js/mermaid-cli -i "{input_path}" -o "{output_path}" -b transparent -w {width} -t {theme}'
        use_shell = True

    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=120,
            shell=use_shell
        )
        if result.returncode != 0:
            raise RuntimeError(f"mermaid rendering failed: {result.stderr}")
        if not os.path.isfile(output_path):
            raise RuntimeError("mermaid rendering produced no output file")
        return output_path
    finally:
        try:
            os.unlink(input_path)
        except OSError:
            pass


def _resolve_svg_css_vars(svg_string):
    """Resolve CSS custom properties in beautiful-mermaid SVGs.

    beautiful-mermaid SVGs use CSS variables (--bg, --fg, etc.) and color-mix()
    which cairosvg cannot interpret, causing black backgrounds and invisible text.
    This function:
      1. Extracts top-level variable values from the inline style attribute
      2. Builds a complete derived-variable map for all --_ prefixed vars
      3. Replaces all var() references with resolved hex colors
      4. Strips any broken color-mix() remnants
    """
    import re

    # ── Step 1: Extract top-level CSS variable definitions ──
    # from style="--bg:#FFFFFF;--fg:#333333;--line:#666666;..."
    var_defs = {}
    style_match = re.search(r'style="([^"]*)"', svg_string[:600])
    if style_match:
        style_str = style_match.group(1)
        for m in re.finditer(r'--([\w-]+)\s*:\s*(#[0-9A-Fa-f]{3,8})', style_str):
            var_defs[m.group(1)] = m.group(2)

    if not var_defs:
        return svg_string  # No CSS vars to resolve

    # ── Step 2: Build derived variable map ──
    # beautiful-mermaid defines --_text, --_node-fill, etc. using color-mix()
    # We map them to sensible resolved colors based on the top-level vars
    bg = var_defs.get('bg', '#FFFFFF')
    fg = var_defs.get('fg', '#333333')
    line = var_defs.get('line', '#666666')
    accent = var_defs.get('accent', '#005EB8')
    muted = var_defs.get('muted', '#999999')
    surface = var_defs.get('surface', '#F5F5F5')
    border = var_defs.get('border', '#CCCCCC')

    derived = {
        '_text': fg,
        '_text-sec': '#555555',
        '_text-muted': '#777777',
        '_text-faint': '#AAAAAA',
        '_line': line,
        '_arrow': accent,
        '_node-fill': surface,
        '_node-stroke': border,
        '_group-fill': bg,
        '_group-hdr': fg,
        '_inner-stroke': border,
        '_key-badge': accent,
        '_fill': surface,
        '_border': border,
        '_accent': accent,
    }

    # Merge: top-level vars take priority, derived fill gaps
    all_vars = {**derived, **var_defs}

    # ── Step 3: Replace var() references everywhere in the SVG ──
    def replace_var(match):
        name = match.group(1)
        if name in all_vars:
            return all_vars[name]
        return match.group(0)

    result = re.sub(r'var\(--([\w-]+)(?:\s*,[^)]*?)?\)', replace_var, svg_string)

    # ── Step 4: Remove the <style> block FIRST ──
    # The <style> block contains CSS variable definitions and color-mix() that
    # cairosvg cannot parse. Remove it before cleaning up inline remnants,
    # because the remnant-cleanup regexes can corrupt the style block's
    # multiline CSS content.
    result = re.sub(r'<style>.*?</style>', '', result, flags=re.DOTALL)

    # ── Step 5: Clean up broken color-mix() remnants in inline attributes ──
    # After var() resolution, inline attribute values may look like:
    #   "#F5F5F5 3%, #FFFFFF))" — leftover from color-mix(in srgb, var(...) 3%, var(...))
    # Extract just the hex color and discard the rest
    result = re.sub(
        r'(#[0-9A-Fa-f]{3,8})\s+\d+%[^";<)]*\)',
        r'\1', result
    )
    # Full color-mix() calls that survived
    result = re.sub(
        r'color-mix\(in\s+srgb\s*,\s*(#[0-9A-Fa-f]{3,8})[^)]*\)',
        r'\1', result
    )
    result = re.sub(r'color-mix\([^)]*\)', '#666666', result)

    # Ensure background is set explicitly
    result = re.sub(r'background:\s*var\(--bg\)', f'background:{bg}', result)

    return result


def _svg_to_png(svg_string, width=800):
    """Convert SVG string to PNG file using available tools.

    Tries cairosvg (Python), then Inkscape CLI. Returns path or None.
    """
    import tempfile

    # Resolve CSS custom properties before conversion (beautiful-mermaid SVGs)
    svg_string = _resolve_svg_css_vars(svg_string)

    # Try cairosvg (pip install cairosvg + Cairo DLLs)
    try:
        # On Windows, add MSYS2 Cairo DLLs directory to PATH before import
        _cairo_dll_dir = os.path.join(os.path.expanduser("~"), "cairo_dlls")
        if os.path.isdir(_cairo_dll_dir):
            if _cairo_dll_dir not in os.environ.get("PATH", ""):
                os.environ["PATH"] = _cairo_dll_dir + os.pathsep + os.environ.get("PATH", "")
            os.add_dll_directory(_cairo_dll_dir)
        import cairosvg
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            cairosvg.svg2png(bytestring=svg_string.encode('utf-8'),
                             write_to=f.name, output_width=width)
            return f.name
    except ImportError:
        pass

    # Try Inkscape CLI
    import subprocess
    import shutil
    if shutil.which("inkscape"):
        with tempfile.NamedTemporaryFile(mode='w', suffix='.svg', delete=False,
                                          encoding='utf-8') as sf:
            sf.write(svg_string)
            svg_path = sf.name
        png_path = svg_path.replace('.svg', '.png')
        try:
            subprocess.run(
                ['inkscape', svg_path, f'--export-filename={png_path}',
                 f'--export-width={width}'],
                capture_output=True, timeout=15
            )
            if os.path.isfile(png_path):
                return png_path
        finally:
            try:
                os.unlink(svg_path)
            except OSError:
                pass

    return None


def add_mermaid_diagram(doc, mermaid_code, width_inches=6.0, caption=None,
                        theme="corporate", mode="auto"):
    """Render mermaid code and insert in document.

    Rendering priority (mode="auto"):
      1. beautiful-mermaid SVG → convert to PNG → embed as image
      2. mmdc CLI → PNG → embed as image
      3. beautiful-mermaid ASCII → embed as monospace code block
      4. Raw mermaid source as code block

    Args:
        doc: Document object
        mermaid_code: Mermaid diagram source
        width_inches: Image width in the document
        caption: Optional caption below the diagram
        theme: Theme name or dict (for beautiful-mermaid) or mmdc theme string
        mode: "auto" (try all), "svg" (beautiful-mermaid only),
              "ascii" (box-drawing text), "mmdc" (legacy CLI), "code" (raw source)

    Returns:
        str: Rendering method used ("svg", "mmdc", "ascii", "code")
    """
    mermaid_code = mermaid_code.strip()

    # Mode: ASCII art (always available if beautiful-mermaid is installed)
    if mode == "ascii":
        ascii_art = render_mermaid_ascii(mermaid_code)
        if ascii_art:
            add_code_block(doc, ascii_art, language="diagram")
            if caption:
                _add_caption(doc, caption)
            return "ascii"
        # Fall through to code block
        mode = "code"

    # Mode: SVG via beautiful-mermaid
    if mode in ("auto", "svg"):
        svg = render_mermaid_svg(mermaid_code, theme=theme)
        if svg:
            png_path = _svg_to_png(svg)
            if png_path:
                try:
                    doc.add_picture(png_path, width=Inches(width_inches))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if caption:
                        _add_caption(doc, caption)
                    return "svg"
                finally:
                    try:
                        os.unlink(png_path)
                    except OSError:
                        pass

            # SVG available but no PNG converter — save SVG and embed as image
            # (Word doesn't natively render SVG in all versions, so try ASCII)
            if mode == "svg":
                # If specifically requested SVG but can't convert, try ASCII
                ascii_art = render_mermaid_ascii(mermaid_code)
                if ascii_art:
                    add_code_block(doc, ascii_art, language="diagram")
                    if caption:
                        _add_caption(doc, caption)
                    return "ascii"

    # Mode: mmdc/npx CLI fallback
    if mode in ("auto", "mmdc"):
        if _check_mmdc() is not None:
            mmdc_theme = theme if isinstance(theme, str) else "default"
            img_path = render_mermaid(mermaid_code, theme=mmdc_theme)
            try:
                doc.add_picture(img_path, width=Inches(width_inches))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    _add_caption(doc, caption)
                return "mmdc"
            finally:
                try:
                    os.unlink(img_path)
                except OSError:
                    pass

    # Mode: ASCII fallback (auto only)
    if mode == "auto":
        ascii_art = render_mermaid_ascii(mermaid_code)
        if ascii_art:
            add_code_block(doc, ascii_art, language="diagram")
            if caption:
                _add_caption(doc, caption)
            return "ascii"

    # Final fallback: raw code block
    doc.add_paragraph("[Diagram — install beautiful-mermaid or mermaid-cli to render]")
    add_code_block(doc, mermaid_code, language="mermaid")
    if caption:
        _add_caption(doc, caption)
    return "code"


def _add_caption(doc, caption):
    """Add a centered italic caption paragraph."""
    cap_para = doc.add_paragraph(caption)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap_para.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string("666666")


# ═══════════════════════════════════════════════════════════════════════════
# CONTENT-AWARE TABLE COLUMN SIZING (borrowed from document-format-skills)
# ═══════════════════════════════════════════════════════════════════════════

def _text_weight(text):
    """Calculate visual width of text. CJK characters count as 2, ASCII as 1."""
    w = 0
    for ch in str(text):
        if '\u4e00' <= ch <= '\u9fff' or '\u3000' <= ch <= '\u303f':
            w += 2
        else:
            w += 1
    return w


def _normalize_pcts(raw_pcts, col_min_pct=8, col_max_pct=50):
    """Clamp column percentages to [min, max] and normalize to sum to 100."""
    clamped = [max(col_min_pct, min(col_max_pct, p)) for p in raw_pcts]
    total = sum(clamped) or 1
    return [round(p / total * 100, 1) for p in clamped]


def auto_size_table_columns(table, usable_width_inches=6.5,
                            col_min_pct=8, col_max_pct=50):
    """Set table column widths based on content character weight.

    Analyzes the maximum text width in each column (using _text_weight),
    converts to percentages, clamps to [min_pct, max_pct], and applies
    widths to both the table grid and individual cells.

    Args:
        table: python-docx Table object
        usable_width_inches: Page width minus margins (default 6.5")
        col_min_pct: Minimum column width as percentage (default 8%)
        col_max_pct: Maximum column width as percentage (default 50%)
    """
    if not table.rows:
        return

    ncols = len(table.columns)
    col_weights = [0] * ncols

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < ncols:
                w = _text_weight(cell.text)
                col_weights[i] = max(col_weights[i], w)

    total_weight = sum(col_weights) or 1
    raw_pcts = [w / total_weight * 100 for w in col_weights]
    pcts = _normalize_pcts(raw_pcts, col_min_pct, col_max_pct)

    # Apply widths to cells
    for i, col in enumerate(table.columns):
        width = Inches(usable_width_inches * pcts[i] / 100)
        for cell in col.cells:
            cell.width = width

    # Also set table grid column widths via OXML for full compatibility
    tblGrid = table._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        gridCols = tblGrid.findall(qn('w:gridCol'))
        for i, gc in enumerate(gridCols):
            if i < ncols:
                gc.set(qn('w:w'), str(int(Emu(Inches(usable_width_inches * pcts[i] / 100)))))


# ═══════════════════════════════════════════════════════════════════════════
# DUAL CJK/LATIN FONT SUPPORT (borrowed from document-format-skills)
# ═══════════════════════════════════════════════════════════════════════════

def set_dual_fonts(run, font_latin="Calibri", font_cjk=None):
    """Set separate fonts for Latin and CJK (East Asian) text on a run.

    python-docx's run.font.name only sets the Latin font. This function
    also sets the w:eastAsia attribute via OXML for proper CJK rendering.

    Args:
        run: python-docx Run object
        font_latin: Font for ASCII/Latin text (e.g., "Calibri", "Times New Roman")
        font_cjk: Font for CJK text (e.g., "SimSun", "Microsoft YaHei").
                  If None, only Latin font is set.
    """
    run.font.name = font_latin
    if font_cjk is None:
        return

    from docx.oxml import OxmlElement
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cjk)
    rFonts.set(qn('w:ascii'), font_latin)
    rFonts.set(qn('w:hAnsi'), font_latin)
    rFonts.set(qn('w:cs'), font_latin)


# ═══════════════════════════════════════════════════════════════════════════
# HIGH-LEVEL DIAGRAM HELPERS
# Generate Mermaid syntax from structured data and render into the document.
# Inspired by Text2Diagram (Mermaid-based) and DeepDiagram (multi-agent).
# ═══════════════════════════════════════════════════════════════════════════

def add_flowchart(doc, steps, direction="TD", width_inches=6.0,
                  caption=None, theme="default", subgraphs=None):
    """Generate and embed a flowchart from structured step data.

    Args:
        doc: Document object
        steps: list of dicts with keys:
            - id: unique node id (e.g., "A", "B")
            - label: display text
            - type: "process" (rectangle), "decision" (diamond),
                    "start" (rounded), "end" (rounded), "io" (parallelogram)
        direction: "TD" (top-down), "LR" (left-right), "BT", "RL"
        width_inches: Image width
        caption: Optional caption
        theme: Mermaid theme
        subgraphs: Optional list of dicts with keys:
            - title: subgraph label
            - nodes: list of node ids belonging to this subgraph

    Example:
        steps = [
            {"id": "A", "label": "Start", "type": "start"},
            {"id": "B", "label": "Process Data", "type": "process", "next": ["C"]},
            {"id": "C", "label": "Valid?", "type": "decision",
             "next": [("D", "Yes"), ("E", "No")]},
            {"id": "D", "label": "Save", "type": "process", "next": ["F"]},
            {"id": "E", "label": "Retry", "type": "process", "next": ["B"]},
            {"id": "F", "label": "Done", "type": "end"},
        ]
        add_flowchart(doc, steps)
    """
    shape_map = {
        "process":  ("[", "]"),
        "decision": ("{", "}"),
        "start":    ("([", "])"),
        "end":      ("([", "])"),
        "io":       ("[/", "/]"),
        "database": ("[(", ")]"),
    }

    lines = [f"graph {direction}"]

    # Subgraph groupings
    node_to_subgraph = {}
    if subgraphs:
        for sg in subgraphs:
            for nid in sg.get("nodes", []):
                node_to_subgraph[nid] = sg["title"]

    # Collect nodes by subgraph
    sg_nodes = {}
    ungrouped = []
    for step in steps:
        sg_title = node_to_subgraph.get(step["id"])
        if sg_title:
            sg_nodes.setdefault(sg_title, []).append(step)
        else:
            ungrouped.append(step)

    def _node_def(step):
        stype = step.get("type", "process")
        l, r = shape_map.get(stype, ("[", "]"))
        label = step["label"].replace('"', "'")
        return f'    {step["id"]}{l}"{label}"{r}'

    # Emit subgraphs
    if subgraphs:
        for sg in subgraphs:
            lines.append(f'    subgraph {sg["title"]}')
            for step in sg_nodes.get(sg["title"], []):
                lines.append(_node_def(step))
            lines.append("    end")

    # Emit ungrouped nodes
    for step in ungrouped:
        lines.append(_node_def(step))

    # Emit edges
    for step in steps:
        for target in step.get("next", []):
            if isinstance(target, tuple):
                tid, label = target
                lines.append(f'    {step["id"]} -->|"{label}"| {tid}')
            else:
                lines.append(f'    {step["id"]} --> {target}')

    mermaid_code = "\n".join(lines)
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


def add_gantt_chart(doc, title, sections, width_inches=6.5,
                    caption=None, theme="default", date_format="YYYY-MM-DD"):
    """Generate and embed a Gantt chart (like Text2Diagram's example).

    Args:
        doc: Document object
        title: Chart title
        sections: list of dicts with keys:
            - name: section name (creates colored swim-lane bands)
            - tasks: list of task dicts with keys:
                - name: task label
                - start: start date string (e.g., "2026-01-15")
                - duration: duration string (e.g., "30d", "2w")
                  OR end: end date string
                - status: "done", "active", "crit" (critical), or None
                - after: task id this depends on (for dependency chains)
                - id: optional task id for dependency references

    Example:
        sections = [
            {"name": "Phase 1", "tasks": [
                {"name": "Define Scope", "id": "t1", "start": "2026-01-01", "duration": "14d", "status": "done"},
                {"name": "Gather Reqs", "id": "t2", "start": "2026-01-10", "duration": "21d"},
            ]},
            {"name": "Phase 2", "tasks": [
                {"name": "Develop", "id": "t3", "after": "t2", "duration": "60d", "status": "active"},
            ]},
        ]
        add_gantt_chart(doc, "Project Timeline", sections)
    """
    lines = [
        "gantt",
        f"    title {title}",
        f"    dateFormat {date_format}",
    ]

    for section in sections:
        lines.append(f"    section {section['name']}")
        for task in section.get("tasks", []):
            parts = [f"    {task['name']}"]

            # Status markers
            status = task.get("status", "")
            markers = []
            if status == "done":
                markers.append("done")
            elif status == "active":
                markers.append("active")
            elif status == "crit":
                markers.append("crit")

            # Task ID
            if task.get("id"):
                markers.append(task["id"])

            if markers:
                parts.append(" :" + ", ".join(markers) + ",")
            else:
                parts.append(" :")

            # Timing
            if task.get("after"):
                parts.append(f" after {task['after']},")
            elif task.get("start"):
                parts.append(f" {task['start']},")

            if task.get("duration"):
                parts.append(f" {task['duration']}")
            elif task.get("end"):
                parts.append(f" {task['end']}")

            lines.append("".join(parts))

    mermaid_code = "\n".join(lines)
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


def add_sequence_diagram(doc, participants, messages, width_inches=6.0,
                          caption=None, theme="default", autonumber=True):
    """Generate and embed a sequence diagram.

    Args:
        doc: Document object
        participants: list of dicts with keys:
            - name: participant identifier
            - label: display label (optional, defaults to name)
            - type: "participant" (box) or "actor" (stick figure)
        messages: list of dicts with keys:
            - from_: source participant name
            - to: target participant name
            - text: message label
            - type: "solid" (->>) or "dashed" (-->>), default "solid"
            - activate: True to activate target (optional)
            - deactivate: True to deactivate source (optional)
            - note: optional note text (placed over the 'to' participant)

    Example:
        participants = [
            {"name": "User", "type": "actor"},
            {"name": "API", "label": "REST API"},
            {"name": "DB", "label": "Database"},
        ]
        messages = [
            {"from_": "User", "to": "API", "text": "POST /order", "activate": True},
            {"from_": "API", "to": "DB", "text": "INSERT order", "activate": True},
            {"from_": "DB", "to": "API", "text": "OK", "type": "dashed", "deactivate": True},
            {"from_": "API", "to": "User", "text": "201 Created", "type": "dashed", "deactivate": True},
        ]
        add_sequence_diagram(doc, participants, messages)
    """
    lines = ["sequenceDiagram"]
    if autonumber:
        lines.append("    autonumber")

    for p in participants:
        ptype = p.get("type", "participant")
        label = p.get("label", p["name"])
        if label != p["name"]:
            lines.append(f"    {ptype} {p['name']} as {label}")
        else:
            lines.append(f"    {ptype} {p['name']}")

    for msg in messages:
        arrow = "->>" if msg.get("type", "solid") == "solid" else "-->>"
        line = f"    {msg['from_']}{arrow}{msg['to']}: {msg['text']}"
        lines.append(line)

        if msg.get("activate"):
            lines.append(f"    activate {msg['to']}")
        if msg.get("deactivate"):
            lines.append(f"    deactivate {msg['from_']}")
        if msg.get("note"):
            lines.append(f"    Note over {msg['to']}: {msg['note']}")

    mermaid_code = "\n".join(lines)
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


def add_er_diagram(doc, entities, width_inches=6.0,
                    caption=None, theme="default"):
    """Generate and embed an ER diagram with Crow's foot notation.

    Args:
        doc: Document object
        entities: list of dicts with keys:
            - name: entity name
            - attributes: list of dicts with {name, type, pk (bool), fk (bool)}
            - relations: list of dicts with {to, label, from_card, to_card}
              cardinality: "||" (exactly one), "o|" (zero or one),
                          "}|" (one or more), "}o" (zero or more)

    Example:
        entities = [
            {"name": "Customer", "attributes": [
                {"name": "id", "type": "int", "pk": True},
                {"name": "name", "type": "string"},
            ], "relations": [
                {"to": "Order", "label": "places", "from_card": "||", "to_card": "}o"},
            ]},
            {"name": "Order", "attributes": [
                {"name": "id", "type": "int", "pk": True},
                {"name": "customer_id", "type": "int", "fk": True},
            ]},
        ]
    """
    lines = ["erDiagram"]

    for entity in entities:
        lines.append(f"    {entity['name']} {{")
        for attr in entity.get("attributes", []):
            constraint = ""
            if attr.get("pk"):
                constraint = " PK"
            elif attr.get("fk"):
                constraint = " FK"
            lines.append(f"        {attr['type']} {attr['name']}{constraint}")
        lines.append("    }")

        for rel in entity.get("relations", []):
            fc = rel.get("from_card", "||")
            tc = rel.get("to_card", "}o")
            label = rel.get("label", "has")
            lines.append(f'    {entity["name"]} {fc}--{tc} {rel["to"]} : "{label}"')

    mermaid_code = "\n".join(lines)
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


def add_timeline(doc, title, events, width_inches=6.0,
                  caption=None, theme="default"):
    """Generate and embed a timeline diagram.

    Args:
        doc: Document object
        title: Timeline title
        events: list of dicts with keys:
            - period: time period label (e.g., "Q1 2026", "January")
            - items: list of event description strings

    Example:
        events = [
            {"period": "Q1 2026", "items": ["Kickoff", "Requirements"]},
            {"period": "Q2 2026", "items": ["Development", "Testing"]},
            {"period": "Q3 2026", "items": ["Launch"]},
        ]
        add_timeline(doc, "Project Milestones", events)
    """
    lines = ["timeline", f"    title {title}"]
    for event in events:
        lines.append(f"    {event['period']}")
        for item in event.get("items", []):
            lines.append(f"        : {item}")

    mermaid_code = "\n".join(lines)
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


def add_mindmap(doc, root, width_inches=6.0, caption=None, theme="default"):
    """Generate and embed a mind map diagram.

    Args:
        doc: Document object
        root: Nested dict representing the mind map tree:
            {"label": "Root Topic", "children": [
                {"label": "Branch 1", "children": [
                    {"label": "Leaf 1a"},
                    {"label": "Leaf 1b"},
                ]},
                {"label": "Branch 2"},
            ]}
        width_inches: Image width
        caption: Optional caption
        theme: Mermaid theme

    Example:
        root = {
            "label": "AI Strategy",
            "children": [
                {"label": "Infrastructure", "children": [
                    {"label": "Azure AI Foundry"},
                    {"label": "Databricks"},
                ]},
                {"label": "Use Cases", "children": [
                    {"label": "RAG / Search"},
                    {"label": "Predictive ML"},
                ]},
            ],
        }
        add_mindmap(doc, root)
    """
    lines = ["mindmap"]

    def _emit(node, depth=1):
        indent = "    " * depth
        lines.append(f"{indent}{node['label']}")
        for child in node.get("children", []):
            _emit(child, depth + 1)

    _emit(root)
    mermaid_code = "\n".join(lines)
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


def add_raw_mermaid(doc, mermaid_code, width_inches=6.0,
                     caption=None, theme="default"):
    """Render raw Mermaid syntax and embed in document.

    Use this when you have pre-written Mermaid code (e.g., from an LLM
    or copied from Text2Diagram/DeepDiagram). Supports all Mermaid diagram
    types including flowchart, sequence, gantt, class, state, ER, pie,
    timeline, mindmap, gitgraph, and more.

    Args:
        doc: Document object
        mermaid_code: Raw Mermaid syntax string
        width_inches: Image width
        caption: Optional caption
        theme: Mermaid theme
    """
    return add_mermaid_diagram(doc, mermaid_code, width_inches=width_inches,
                                caption=caption, theme=theme)


# ═══════════════════════════════════════════════════════════════════════════
# D2 DIAGRAM SUPPORT
# D2 (https://d2lang.com) produces architecture diagrams with icons,
# dashed regions, and cloud-provider shapes. Requires `d2` CLI binary.
# ═══════════════════════════════════════════════════════════════════════════

def _check_d2():
    """Check if d2 CLI is available (PATH or known install location)."""
    import shutil
    if shutil.which("d2") is not None:
        return True
    # Check known install location on this machine
    known = r"<USER_HOME>/tools\d2\d2-v0.7.1\bin\d2.exe"
    if os.path.isfile(known):
        return True
    return False


def _get_d2_cmd():
    """Return the d2 command (full path if not on PATH)."""
    import shutil
    if shutil.which("d2") is not None:
        return "d2"
    known = r"<USER_HOME>/tools\d2\d2-v0.7.1\bin\d2.exe"
    if os.path.isfile(known):
        return known
    return "d2"


def render_d2(d2_code, output_format="png", theme=0, layout="elk"):
    """Render D2 code to an image file.

    Requires: d2 CLI (install from https://d2lang.com or `curl -fsSL https://d2lang.com/install.sh | sh`)

    Args:
        d2_code: D2 diagram source code
        output_format: "png" or "svg"
        theme: D2 theme number (0=default, 1=neutral, 3=vanilla, 100=dark)
        layout: Layout engine ("dagre", "elk", "tala")

    Returns:
        Path to rendered image (caller must delete after use).
    """
    import subprocess
    import tempfile

    if not _check_d2():
        raise RuntimeError(
            "d2 not found. Install from https://d2lang.com"
        )

    with tempfile.NamedTemporaryFile(mode='w', suffix='.d2', delete=False,
                                      encoding='utf-8') as f:
        f.write(d2_code)
        input_path = f.name

    output_path = input_path.replace('.d2', f'.{output_format}')

    try:
        d2_cmd = _get_d2_cmd()
        result = subprocess.run(
            [d2_cmd, '--theme', str(theme), '--layout', layout,
             input_path, output_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            raise RuntimeError(f"d2 rendering failed: {result.stderr}")
        if not os.path.isfile(output_path):
            raise RuntimeError("d2 rendering produced no output file")
        return output_path
    finally:
        try:
            os.unlink(input_path)
        except OSError:
            pass


def add_d2_diagram(doc, d2_code, width_inches=6.0, caption=None,
                    theme=0, layout="elk"):
    """Render D2 code and embed in document.

    D2 excels at architecture diagrams, cloud infrastructure,
    and diagrams with icons. Falls back to code block if d2 CLI
    is not installed.

    Args:
        doc: Document object
        d2_code: D2 diagram source code
        width_inches: Image width
        caption: Optional caption
        theme: D2 theme (0=default, 1=neutral, 3=vanilla)
        layout: Layout engine ("dagre", "elk", "tala")

    Returns:
        str: "d2" if rendered, "code" if fallback
    """
    if not _check_d2():
        doc.add_paragraph("[D2 Diagram — install d2 CLI from https://d2lang.com to render]")
        add_code_block(doc, d2_code, language="d2")
        if caption:
            _add_caption(doc, caption)
        return "code"

    img_path = render_d2(d2_code, theme=theme, layout=layout)
    try:
        doc.add_picture(img_path, width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            _add_caption(doc, caption)
        return "d2"
    finally:
        try:
            os.unlink(img_path)
        except OSError:
            pass


# ═══════════════════════════════════════════════════════════════════════════
# MATPLOTLIB-BASED CHARTS (for data-driven visuals in DOCX)
# Inspired by docx-report's matplotlib integration pattern.
# ═══════════════════════════════════════════════════════════════════════════

def add_matplotlib_chart(doc, plot_func, width_inches=5.5, caption=None, dpi=150):
    """Embed a matplotlib figure in the document.

    Args:
        doc: Document object
        plot_func: A callable that takes (fig, ax) and draws the chart.
                   The function should NOT call plt.show() or fig.savefig().
        width_inches: Image width in the document
        caption: Optional caption
        dpi: Resolution (default 150 for print quality)

    Example:
        def my_chart(fig, ax):
            ax.bar(["Q1", "Q2", "Q3", "Q4"], [100, 150, 130, 180])
            ax.set_title("Quarterly Revenue")
            ax.set_ylabel("$M")

        add_matplotlib_chart(doc, my_chart, caption="Figure 1: Revenue by Quarter")
    """
    import tempfile
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
    except ImportError:
        doc.add_paragraph("[Chart — install matplotlib to render]")
        if caption:
            _add_caption(doc, caption)
        return "code"

    fig, ax = plt.subplots(figsize=(width_inches, width_inches * 0.6))
    plot_func(fig, ax)
    fig.tight_layout()

    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
        fig.savefig(f.name, dpi=dpi, bbox_inches='tight',
                    facecolor='white', edgecolor='none')
        img_path = f.name
    plt.close(fig)

    try:
        doc.add_picture(img_path, width=Inches(width_inches))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            _add_caption(doc, caption)
        return "matplotlib"
    finally:
        try:
            os.unlink(img_path)
        except OSError:
            pass


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT UTILITIES — Watermarks, Bookmarks, Comments, Footnotes, Sections
# ═══════════════════════════════════════════════════════════════════════════

def add_watermark(doc, text="DRAFT", color="C0C0C0", font_size=72, rotation=-45):
    """Add a diagonal text watermark to every section of the document.

    Uses VML (Vector Markup Language) shapes in the header, which is how
    Word natively implements watermarks.

    Args:
        doc: Document object
        text: Watermark text (e.g., "DRAFT", "CONFIDENTIAL")
        color: Hex color without '#' (default light gray)
        font_size: Font size in points
        rotation: Rotation angle in degrees (negative = counter-clockwise)

    Example:
        doc = create_document(preset="executive")
        doc.add_paragraph("This is a draft document.")
        add_watermark(doc, text="DRAFT", color="C0C0C0")
        doc.save("draft_report.docx")
    """
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        run = p.add_run()
        pict_xml = (
            '<w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:v="urn:schemas-microsoft-com:vml"'
            ' xmlns:o="urn:schemas-microsoft-com:office:office">'
            '<v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136"'
            ' path="m@7,l@8,m@5,21600l@6,21600e">'
            '<v:formulas>'
            '<v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>'
            '<v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/>'
            '<v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/>'
            '<v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @4"/>'
            '<v:f eqn="if @0 @2 21600"/><v:f eqn="mid @5 @6"/>'
            '<v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/>'
            '<v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>'
            '</v:formulas>'
            '<v:path textpathok="t" o:connecttype="custom"/>'
            '<v:textpath on="t" fitshape="t"/>'
            '<o:lock v:ext="edit" text="t" shapetype="t"/>'
            '</v:shapetype>'
            f'<v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049"'
            f' type="#_x0000_t136"'
            f' style="position:absolute;margin-left:0;margin-top:0;'
            f'width:527.85pt;height:131.95pt;rotation:{rotation};'
            f'z-index:-251658752;'
            f'mso-position-horizontal:center;'
            f'mso-position-horizontal-relative:margin;'
            f'mso-position-vertical:center;'
            f'mso-position-vertical-relative:margin"'
            f' o:allowincell="f" fillcolor="#{color}" stroked="f">'
            f'<v:fill opacity=".5"/>'
            f'<v:textpath style="font-family:&amp;quot;Calibri&amp;quot;;'
            f'font-size:{font_size}pt" string="{text}"/>'
            '</v:shape>'
            '</w:pict>'
        )
        try:
            pict_elem = parse_xml(pict_xml)
            run._element.append(pict_elem)
        except Exception as e:
            import warnings
            warnings.warn(f"add_watermark: failed to insert VML shape — {e}")


def add_bookmark(paragraph, name):
    """Insert a bookmark at the given paragraph for cross-referencing.

    Creates a bookmarkStart/bookmarkEnd pair wrapping the paragraph content.
    The bookmark can be referenced from hyperlinks or TOC fields.

    Args:
        paragraph: python-docx Paragraph object
        name: Bookmark name (must be unique within the document, no spaces)

    Example:
        p = doc.add_paragraph("Important Section")
        add_bookmark(p, "important_section")
    """
    bm_id = str(abs(hash(name)) % 99999)
    bm_start = parse_xml(
        f'<w:bookmarkStart {nsdecls("w")} w:id="{bm_id}" w:name="{name}"/>')
    bm_end = parse_xml(
        f'<w:bookmarkEnd {nsdecls("w")} w:id="{bm_id}"/>')
    paragraph._element.insert(0, bm_start)
    paragraph._element.append(bm_end)


def add_comment(doc, paragraph, comment_text, author="Author", initials="A",
                date_str=None):
    """Add a comment annotation to a paragraph (experimental).

    Creates the comments XML part if it does not exist, appends the comment,
    and inserts commentRangeStart/End markers in the paragraph.

    Note:
        This is an OXML-level manipulation that works reliably with documents
        created by this module. Complex documents with existing comments may
        need manual adjustment.

    Args:
        doc: Document object
        paragraph: Target paragraph to annotate
        comment_text: The comment body text
        author: Comment author name
        initials: Author initials shown in the comment balloon
        date_str: ISO datetime string (defaults to now)

    Example:
        p = doc.add_paragraph("Revenue grew 15% YoY.")
        add_comment(doc, p, "Verify with Q4 financials", author="Reviewer")
    """
    from datetime import datetime
    if date_str is None:
        date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    part = doc.part
    comments_reltype = (
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
    )

    # Locate existing comments part
    comments_part = None
    for rel in part.rels.values():
        if rel.reltype == comments_reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        comments_xml = (
            '<w:comments'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        )
        comments_part = Part(
            PackURI('/word/comments.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
            comments_xml.encode('utf-8'),
            part.package
        )
        part.relate_to(comments_part, comments_reltype)

    # Parse existing comments and determine next ID
    comments_elem = etree.fromstring(comments_part.blob)
    existing_ids = [
        int(c.get(qn('w:id'), '0'))
        for c in comments_elem.findall(qn('w:comment'))
    ]
    comment_id = str(max(existing_ids, default=-1) + 1)

    # Append new comment element
    comment_elem = parse_xml(
        f'<w:comment {nsdecls("w")} w:id="{comment_id}"'
        f' w:author="{author}" w:initials="{initials}" w:date="{date_str}">'
        f'<w:p><w:r><w:t>{comment_text}</w:t></w:r></w:p>'
        f'</w:comment>'
    )
    comments_elem.append(comment_elem)
    comments_part._blob = etree.tostring(
        comments_elem, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # Insert comment range markers and reference in paragraph
    p = paragraph._element
    range_start = parse_xml(
        f'<w:commentRangeStart {nsdecls("w")} w:id="{comment_id}"/>')
    range_end = parse_xml(
        f'<w:commentRangeEnd {nsdecls("w")} w:id="{comment_id}"/>')
    ref_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
        f'<w:commentReference w:id="{comment_id}"/>'
        f'</w:r>'
    )
    p.insert(0, range_start)
    p.append(range_end)
    p.append(ref_run)


def add_footnote(paragraph, text):
    """Add a footnote to a paragraph (experimental).

    Creates the footnotes XML part (with separators) if needed, appends the
    footnote content, and inserts a superscript footnote reference at the
    end of the paragraph.

    Note:
        This is an OXML-level manipulation that works reliably with documents
        created by this module. Documents opened from templates with existing
        footnotes may need manual verification.

    Args:
        paragraph: python-docx Paragraph object
        text: Footnote body text

    Example:
        p = doc.add_paragraph("The market expanded significantly.")
        add_footnote(p, "Based on 2025 Q3 earnings reports.")
    """
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    doc_part = paragraph.part
    footnotes_reltype = (
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
    )

    # Locate existing footnotes part
    footnotes_part = None
    for rel in doc_part.rels.values():
        if 'footnotes' in rel.reltype:
            footnotes_part = rel.target_part
            break

    if footnotes_part is None:
        fn_xml = (
            '<w:footnotes'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
            '</w:footnotes>'
        )
        footnotes_part = Part(
            PackURI('/word/footnotes.xml'),
            'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml',
            fn_xml.encode('utf-8'),
            doc_part.package
        )
        doc_part.relate_to(footnotes_part, footnotes_reltype)

    # Parse existing footnotes and determine next ID
    footnotes_elem = etree.fromstring(footnotes_part.blob)
    existing_ids = [
        int(fn.get(qn('w:id'), '0'))
        for fn in footnotes_elem.findall(qn('w:footnote'))
    ]
    fn_id = str(max(existing_ids, default=0) + 1)

    # Append footnote content
    fn_elem = parse_xml(
        f'<w:footnote {nsdecls("w")} w:id="{fn_id}">'
        f'<w:p><w:pPr><w:pStyle w:val="FootnoteText"/></w:pPr>'
        f'<w:r><w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteRef/></w:r>'
        f'<w:r><w:t xml:space="preserve"> {text}</w:t></w:r></w:p>'
        f'</w:footnote>'
    )
    footnotes_elem.append(fn_elem)
    footnotes_part._blob = etree.tostring(
        footnotes_elem, xml_declaration=True, encoding='UTF-8', standalone=True
    )

    # Insert footnote reference in the paragraph
    ref_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rStyle w:val="FootnoteReference"/></w:rPr>'
        f'<w:footnoteReference w:id="{fn_id}"/>'
        f'</w:r>'
    )
    paragraph._element.append(ref_run)


def add_section_break(doc, break_type="nextPage"):
    """Add a section break to the document.

    Args:
        doc: Document object
        break_type: One of "nextPage", "continuous", "evenPage", "oddPage"

    Returns:
        The new Section object.

    Example:
        doc = create_document(preset="executive")
        doc.add_paragraph("Page 1 content")
        new_section = add_section_break(doc, "continuous")
        doc.add_paragraph("Continues on same page")
    """
    from docx.enum.section import WD_SECTION_START
    type_map = {
        "nextPage":   WD_SECTION_START.NEW_PAGE,
        "continuous": WD_SECTION_START.CONTINUOUS,
        "evenPage":   WD_SECTION_START.EVEN_PAGE,
        "oddPage":    WD_SECTION_START.ODD_PAGE,
    }
    start_type = type_map.get(break_type, WD_SECTION_START.NEW_PAGE)
    new_section = doc.add_section(start_type)
    return new_section


def set_paragraph_flow(paragraph, keep_next=None, keep_together=None,
                       page_break_before=None, widow_control=None):
    """Set paragraph flow-control properties via OXML.

    These properties control how Word handles page breaks around a paragraph.

    Args:
        paragraph: python-docx Paragraph object
        keep_next: If True, keep this paragraph on the same page as the next
        keep_together: If True, prevent page break within this paragraph
        page_break_before: If True, force a page break before this paragraph
        widow_control: If True, prevent orphan/widow lines

    Example:
        heading = doc.add_heading("Section Title", level=2)
        set_paragraph_flow(heading, keep_next=True, keep_together=True)
    """
    pPr = paragraph._element.get_or_add_pPr()
    _props = {
        'w:keepNext':        keep_next,
        'w:keepLines':       keep_together,
        'w:pageBreakBefore': page_break_before,
        'w:widowControl':    widow_control,
    }
    for tag, value in _props.items():
        if value is None:
            continue
        elem = pPr.find(qn(tag))
        if value:
            if elem is None:
                etree.SubElement(pPr, qn(tag))
        else:
            if elem is not None:
                pPr.remove(elem)


def find_and_replace(doc, old_text, new_text, match_case=True):
    """Search-and-replace text across all paragraphs and table cells.

    Handles text that is split across multiple runs by joining run texts
    and redistributing the result.

    Args:
        doc: Document object
        old_text: Text to find
        new_text: Replacement text
        match_case: If False, perform case-insensitive matching

    Returns:
        int: Number of paragraphs where a replacement was made.

    Example:
        doc = Document("template.docx")
        n = find_and_replace(doc, "{{CLIENT_NAME}}", "Acme Corp")
        print(f"Replaced in {n} locations")
    """
    count = 0
    for paragraph in doc.paragraphs:
        count += _replace_in_paragraph(paragraph, old_text, new_text, match_case)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += _replace_in_paragraph(
                        paragraph, old_text, new_text, match_case
                    )
    return count


def _replace_in_paragraph(paragraph, old_text, new_text, match_case):
    """Replace text within a single paragraph, handling cross-run spans.

    Args:
        paragraph: python-docx Paragraph object
        old_text: Text to find
        new_text: Replacement text
        match_case: If False, case-insensitive

    Returns:
        int: 1 if a replacement was made, 0 otherwise.
    """
    full_text = paragraph.text
    if not match_case:
        if old_text.lower() not in full_text.lower():
            return 0
    elif old_text not in full_text:
        return 0

    # Fast path: text lives entirely within a single run
    for run in paragraph.runs:
        run_text = run.text if match_case else run.text.lower()
        target = old_text if match_case else old_text.lower()
        if target in run_text:
            if match_case:
                run.text = run.text.replace(old_text, new_text)
            else:
                # Preserve surrounding case by using positional replacement
                idx = run.text.lower().find(old_text.lower())
                run.text = run.text[:idx] + new_text + run.text[idx + len(old_text):]
            return 1

    # Slow path: text spans multiple runs — combine, replace, redistribute
    combined = "".join(run.text for run in paragraph.runs)
    if match_case:
        new_combined = combined.replace(old_text, new_text)
    else:
        # Case-insensitive replacement
        import re as _re
        new_combined = _re.sub(_re.escape(old_text), new_text, combined,
                               flags=_re.IGNORECASE)

    if new_combined != combined and paragraph.runs:
        paragraph.runs[0].text = new_combined
        for run in paragraph.runs[1:]:
            run.text = ""
        return 1
    return 0


def add_paragraph_border(paragraph, style="single", width_pt=1,
                         color="000000", space_pt=1):
    """Add a box border around a paragraph via OXML pBdr element.

    Args:
        paragraph: python-docx Paragraph object
        style: Border style — "single", "double", "dashed", "dotted",
               "thick", "wave", "dashSmallGap", "dotDash", "dotDotDash"
        width_pt: Border width in points (converted to eighths internally)
        color: Hex color without '#'
        space_pt: Space between border and text in points

    Example:
        p = doc.add_paragraph("Important notice — read carefully.")
        add_paragraph_border(p, style="single", width_pt=1.5, color="DC3545")
    """
    pPr = paragraph._element.get_or_add_pPr()

    # Remove any existing paragraph borders
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)

    sz = str(int(width_pt * 8))   # Word uses eighths of a point
    space = str(int(space_pt))

    border_xml = (
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:top w:val="{style}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'<w:left w:val="{style}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'<w:bottom w:val="{style}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'<w:right w:val="{style}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pBdr = parse_xml(border_xml)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════
# MAIN (CLI)
# ═══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="Convert Markdown to beautified DOCX")
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("-o", "--output", help="Output DOCX file")
    parser.add_argument("-p", "--preset", default="executive",
                        choices=list(PRESETS.keys()), help="Document preset")
    parser.add_argument("--title", help="Document title")
    parser.add_argument("--author", help="Author name")
    parser.add_argument("--cover", action="store_true", help="Add cover page")
    parser.add_argument("--no-header-footer", action="store_true",
                        help="Skip header/footer")

    args = parser.parse_args()
    output = args.output or args.input.rsplit(".", 1)[0] + ".docx"

    md_to_docx(
        args.input, output,
        preset=args.preset,
        title=args.title,
        author=args.author,
        cover_page=args.cover,
        header_footer=not args.no_header_footer,
    )
